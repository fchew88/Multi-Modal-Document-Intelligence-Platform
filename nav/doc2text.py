import streamlit as st
import pandas as pd
import pdfplumber
import docx
from openai import OpenAI
import tempfile
import os
import re
import json

# Initialize OpenAI client
try:
    client = OpenAI(api_key=st.secrets["openai_api_key"])
except Exception as e:
    st.error(f"Failed to initialize OpenAI client: {str(e)}")
    client = None

st.title('📄 Smart Document Analyzer')

# File processing functions
def extract_text_from_pdf(file):
    """Extract text from PDF using pdfplumber"""
    with pdfplumber.open(file) as pdf:
        text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    return text

def extract_text_from_csv(file):
    """Extract text content from CSV file"""
    try:
        df = pd.read_csv(file)
        # Convert DataFrame to readable text
        text_content = f"CSV File Contents:\n\n{df.to_string(index=False)}"
        
        # Also extract the data as tables
        tables = [df.values.tolist()]
        
        return text_content, tables
    except Exception as e:
        st.error(f"Failed to process CSV file: {str(e)}")
        return None, None

def extract_text_from_docx(file):
    """Extract text from DOCX using python-docx"""
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_xlsx(file):
    """Extract text content from Excel file"""
    try:
        # Read all sheets from the Excel file
        xls = pd.ExcelFile(file)
        text_content = []
        tables = []
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text_content.append(f"\n\nSheet: {sheet_name}\n\n{df.to_string(index=False)}")
            tables.append(df.values.tolist())
        
        return "\n".join(text_content), tables
    except Exception as e:
        st.error(f"Failed to process Excel file: {str(e)}")
        return None, None

def extract_tables_from_pdf(file):
    """Extract tables from PDF using pdfplumber with better error handling"""
    tables = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            try:
                page_tables = page.extract_tables()
                if page_tables:
                    for table in page_tables:
                        # Clean table data
                        cleaned_table = []
                        for row in table:
                            cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                            cleaned_table.append(cleaned_row)
                        tables.append(cleaned_table)
            except Exception as e:
                st.warning(f"Could not extract tables from page {page.page_number}: {str(e)}")
    return tables

def clean_text(text):
    """Clean extracted text"""
    text = re.sub(r'\s+', ' ', text)  # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)  # Limit consecutive newlines
    return text.strip()

def structure_text(text):
    """AI-powered document structure analysis with regex fallback"""
    structured = []
    
    # First try using AI for structure analysis if client is available
    if client and len(text) > 500:  # Only use AI for substantial documents
        try:
            # Prepare the prompt
            prompt = f"""
            Analyze this document text and identify its logical structure. 
            Return the analysis in JSON format with sections containing:
            - "heading" (the section title if detectable)
            - "content" (the section text)
            - "type" (section type: title, heading, paragraph, list, etc.)
            
            Rules:
            1. Preserve all original text content
            2. Group related paragraphs under headings
            3. Identify section types based on formatting and content
            4. Handle numbered/bulleted lists appropriately
            5. Return valid JSON format
            
            Document Text:
            {text[:15000]}  # Limit to first 15k chars to avoid token limits
            """
            
            # Get AI response
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                response_format={"type": "json_object"}
            )
            
            # Parse the AI response
            ai_structure = json.loads(response.choices[0].message.content)
            
            # Convert AI structure to our format
            for section in ai_structure.get('sections', []):
                content = section.get('content', '').strip()
                if content and len(content.split()) > 5:
                    structured.append({
                        'content': content,
                        'type': section.get('type', 'section'),
                        'length': len(content),
                        'heading': section.get('heading', '')
                    })
            
            if structured:
                return structured
                
        except Exception as e:
            st.warning(f"AI structure analysis failed, using fallback: {str(e)}")
    
    # Fallback to regex-based structuring if AI fails or isn't available
    section_patterns = [
        r'\n\s*\n\s*[A-Z][A-Z0-9 ]+\s*\n\s*\n',  # All-caps headings
        r'\n\s*\d+\.\s+[A-Z][^\n]+\s*\n',         # Numbered headings (1. SECTION)
        r'\n\s*[A-Z][^\n]+\s*\n[-=]+\s*\n',       # Underlined headings
        r'\n\s*[A-Z][^\n]+:\s*\n',                # Headings ending with colon
        r'\n\s*[•■♦●]\s+[A-Z][^\n]+\s*\n',        # Bullet point headings
        r'\n\s*\n\s*[A-Z][^\n]+\s*\n\s*\n',       # Regular headings with spacing
    ]
    
    for pattern in section_patterns:
        sections = re.split(pattern, text)
        if len(sections) > 1:
            for section in sections:
                if len(section.strip()) > 0 and len(section.split()) > 5:
                    structured.append({
                        'content': section.strip(),
                        'type': 'section',
                        'length': len(section.strip()),
                        'heading': ''
                    })
            break
    
    if not structured:
        paragraphs = re.split(r'\n\s*\n', text)
        for para in paragraphs:
            if len(para.strip()) > 0 and len(para.split()) > 5:
                structured.append({
                    'content': para.strip(),
                    'type': 'paragraph', 
                    'length': len(para.strip()),
                    'heading': ''
                })
    
    return structured

# Initialize session state
if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = None
if 'structured_data' not in st.session_state:
    st.session_state.structured_data = None
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None

# Document selection section
option = st.radio("Select document source:", ("Upload a document", "Use sample document"))

if option == "Upload a document":
    uploaded_file = st.file_uploader("Choose a document file", 
                                   type=["pdf", "docx", "txt", "csv", "xlsx"])
else:
    # Load sample document
    sample_path = "data/Doc/GT-FY23-FinS.pdf"
    try:
        if os.path.exists(sample_path):
            with open(sample_path, "rb") as f:
                uploaded_file = f.read()
        else:
            st.error(f"Sample document not found at: {sample_path}")
            uploaded_file = None
    except Exception as e:
        st.error(f"Failed to load sample document: {str(e)}")
        uploaded_file = None

if uploaded_file is not None:
    # Display file info
    if option == "Upload a document":
        file_extension = uploaded_file.name.split('.')[-1].lower()
        st.success(f"Uploaded {file_extension.upper()} file: {uploaded_file.name}")
    else:
        st.success("Sample document loaded: GT-FY23-FinS.pdf")
        file_extension = "pdf"

    if st.button('Process Document'):
        with st.spinner('Processing document...'):
            try:
                with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                    if option == "Upload a document":
                        tmp_file.write(uploaded_file.getvalue())
                    else:
                        tmp_file.write(uploaded_file)
                    tmp_file_path = tmp_file.name
                
                text = ""
                tables = []
                
                if file_extension == 'pdf':
                    text = extract_text_from_pdf(tmp_file_path)
                    tables = extract_tables_from_pdf(tmp_file_path)
                elif file_extension == 'docx':
                    text = extract_text_from_docx(tmp_file_path)
                elif file_extension == 'txt':
                    if option == "Upload a document":
                        text = uploaded_file.read().decode("utf-8")
                    else:
                        with open(tmp_file_path, "r", encoding="utf-8") as f:
                            text = f.read()
                elif file_extension == 'csv':
                    if option == "Upload a document":
                        text, tables = extract_text_from_csv(uploaded_file)
                    else:
                        text, tables = extract_text_from_csv(tmp_file_path)
                    if text:
                        text = clean_text(text)
                        st.session_state.extracted_text = text
                        st.session_state.structured_data = structure_text(text)
                    if tables:
                        st.session_state.tables = tables
                elif file_extension == 'xlsx':
                    if option == "Upload a document":
                        text, tables = extract_text_from_xlsx(uploaded_file)
                    else:
                        text, tables = extract_text_from_xlsx(tmp_file_path)
                    if text:
                        text = clean_text(text)
                        st.session_state.extracted_text = text
                        st.session_state.structured_data = structure_text(text)
                    if tables:
                        st.session_state.tables = tables
                
                # Clean and structure the text
                if text:
                    text = clean_text(text)
                    st.session_state.extracted_text = text
                    st.session_state.structured_data = structure_text(text)
                
                # Process tables
                if tables:
                    st.session_state.tables = tables
                
                os.unlink(tmp_file_path)
                
                st.success("Document processed successfully!")
                
            except Exception as e:
                st.error(f"Document processing failed: {str(e)}")

# Display results
if st.session_state.get('extracted_text'):
    st.subheader("📝 Document Content")
    
    with st.expander("View Full Text"):
        st.text_area("Text Content", 
                    value=st.session_state.extracted_text, 
                    height=300,
                    label_visibility="collapsed")
    
    if st.session_state.get('tables'):
        st.subheader("📊 Extracted Tables")
        for i, table in enumerate(st.session_state.tables[:3]):  # Show first 3 tables
            with st.expander(f"Table {i+1}"):
                try:
                    if len(table) > 1:
                        headers = table[0]
                        data = table[1:]
                        
                        # Generate unique column names
                        seen = {}
                        unique_headers = []
                        for i, h in enumerate(headers):
                            if pd.isna(h) or h == '':
                                unique_headers.append(f"Column_{i+1}")
                            elif h in seen:
                                seen[h] += 1
                                unique_headers.append(f"{h}_{seen[h]}")
                            else:
                                seen[h] = 0
                                unique_headers.append(h)
                        
                        # Create DataFrame and replace NaN with empty strings
                        df = pd.DataFrame(data, columns=unique_headers)
                        df = df.fillna('')  # Replace all NaN values
                        
                        st.dataframe(df)
                    else:
                        st.warning(f"Table {i+1} doesn't have enough rows to display")
                except Exception as e:
                    st.error(f"Could not display table {i+1}: {str(e)}")
                    st.write("Raw table data (first 5 rows):")
                    st.write(table[:5])

# AI Analysis Section
if st.session_state.get('extracted_text') and client:
    st.subheader("🤖 AI Document Analysis")
    
    analysis_type = st.radio("Analysis Type",
                            ["Summarize", "Extract Key Points", "Find Action Items", "Custom Query"],
                            horizontal=True)
    
    query = ""
    if analysis_type == "Custom Query":
        query = st.text_input("Enter your analysis query")
    else:
        query = analysis_type
    
    if st.button("Run Analysis"):
        with st.spinner("Analyzing document..."):
            try:
                prompt = f"""
                Perform this analysis on the document: {query}
                
                Document Content:
                {st.session_state.extracted_text[:10000]}  # Limit to first 10k chars
                
                Instructions:
                1. Be concise and structured
                2. Use bullet points for key information
                3. Highlight important names, dates, and figures
                """
                
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5
                )
                
                st.session_state.analysis_results = response.choices[0].message.content
                
                st.subheader("Analysis Results")
                st.markdown(st.session_state.analysis_results)
                
            except Exception as e:
                st.error(f"Analysis failed: {str(e)}")

# Clear button
if st.button("Clear All"):
    st.session_state.clear()
    st.success("All content cleared!")
    st.rerun()

# Debug view
#with st.expander("Debug: Session State"):
#    st.write(st.session_state)